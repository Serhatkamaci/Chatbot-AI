import os
import shutil
import json
import asyncio
import nest_asyncio
from dotenv import load_dotenv

# Nested event loops'a izin ver
nest_asyncio.apply()
from flask import Flask, request, jsonify, Response, send_from_directory
from flask_cors import CORS
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from docx import Document
from langchain.schema import Document as LangchainDocument
from werkzeug.utils import secure_filename
import uuid
from datetime import datetime

# Konfigürasyon
load_dotenv()
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max dosya boyutu
CORS(app)

google_api_key = os.getenv("GOOGLE_API_KEY")
if not google_api_key:
    raise ValueError("GOOGLE_API_KEY ortam değişkeni ayarlanmamış.")

# Global değişkenler
UPLOAD_FOLDER = 'uploads'
DATA_FOLDER = 'data'
STATIC_FOLDER = 'static'
HISTORY_FILE = "kubik_history.json"
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Klasörleri oluştur
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DATA_FOLDER, exist_ok=True)
os.makedirs(STATIC_FOLDER, exist_ok=True)

# Aktif chatbot'ları saklamak için
active_chatbots = {}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# --- Kübik geçmişini yönetme ---
def load_kubik_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []


def save_kubik_history(history):
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def add_kubik_to_history(kubik_name):
    history = load_kubik_history()
    history = [h for h in history if h != kubik_name]
    history.insert(0, kubik_name)
    if len(history) > 20:
        history = history[:20]
    save_kubik_history(history)
    return history


def remove_kubik_from_history(kubik_name):
    history = load_kubik_history()
    history = [h for h in history if h != kubik_name]
    save_kubik_history(history)
    return history


# --- DOCX dosyasını güvenli açma ve metin alma ---
def load_docx_with_python_docx(file_path):
    try:
        doc = Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip() != ""]
        if not paragraphs:
            print(f"DOCX dosyası içerik olarak boş: {file_path}")
            return []
        full_text = "\n".join(paragraphs)
        return [LangchainDocument(page_content=full_text, metadata={"source": file_path})]
    except Exception as e:
        print(f"DOCX dosyası açılamadı ({file_path}): {e}")
        return []


# --- Belgeleri Yükle ve Böl ---
def load_and_process_documents(folder_path):
    documents = []
    pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".pdf")]
    docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]

    if not pdf_files and not docx_files:
        print(f"'{folder_path}' klasöründe PDF veya DOCX dosyası bulunamadı.")
        return None

    for pdf_file in pdf_files:
        file_path = os.path.join(folder_path, pdf_file)
        print(f"'{pdf_file}' yükleniyor...")
        loader = PyPDFLoader(file_path)
        documents.extend(loader.load())

    for docx_file in docx_files:
        file_path = os.path.join(folder_path, docx_file)
        print(f"'{docx_file}' yükleniyor...")
        documents.extend(load_docx_with_python_docx(file_path))

    if not documents:
        return None

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=300,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    return text_splitter.split_documents(documents)


# --- Vektör Veritabanı Yarat veya Yükle ---
def create_or_load_vector_store(docs, index_path):
    # Embeddings'i thread-safe bir şekilde oluştur
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=google_api_key
    )

    if os.path.exists(index_path):
        print("Vektör veritabanı yükleniyor...")
        try:
            return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        except Exception as e:
            print(f"Vektör veritabanı yüklenemedi: {e}")
            # Yükleme başarısızsa yeniden oluştur
            if docs:
                print("Yeni vektör veritabanı oluşturuluyor...")
                vector_store = FAISS.from_documents(docs, embeddings)
                vector_store.save_local(index_path)
                return vector_store
            return None

    if not docs:
        print("Vektör veritabanı oluşturmak için belge yok.")
        return None

    print("Yeni vektör veritabanı oluşturuluyor...")
    try:
        vector_store = FAISS.from_documents(docs, embeddings)
        vector_store.save_local(index_path)
        return vector_store
    except Exception as e:
        print(f"Vektör veritabanı oluşturulamadı: {e}")
        return None


# --- LLM Wrapper for Thread Safety ---
class ThreadSafeLLM:
    def __init__(self, google_api_key, temperature=0.2):
        self.google_api_key = google_api_key
        self.temperature = temperature
        self._llm = None

    def get_llm(self):
        if self._llm is None:
            self._llm = ChatGoogleGenerativeAI(
                model="gemini-1.5-flash-latest",
                google_api_key=self.google_api_key,
                temperature=self.temperature,
            )
        return self._llm

    def invoke(self, prompt):
        llm = self.get_llm()
        return llm.invoke(prompt)

    def stream(self, prompt):
        llm = self.get_llm()
        return llm.stream(prompt)


# --- Kübik dosya bilgilerini al ---
def get_kubik_files(kubik_name):
    folder_path = os.path.join(DATA_FOLDER, kubik_name)
    if not os.path.exists(folder_path):
        return []

    files = []
    for file in os.listdir(folder_path):
        if file.lower().endswith(('.pdf', '.docx')) and not file.startswith('.'):
            files.append(file)

    return files


# --- API Endpoints ---

@app.route('/api/kubik/history', methods=['GET'])
def get_kubik_history():
    """Kübik geçmişini getir"""
    history = load_kubik_history()
    return jsonify({'success': True, 'history': history})


@app.route('/api/kubik/info/<kubik_name>', methods=['GET'])
def get_kubik_info(kubik_name):
    """Kübik bilgilerini getir"""
    try:
        folder_path = os.path.join(DATA_FOLDER, kubik_name)
        if not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Kübik bulunamadı'})

        files = get_kubik_files(kubik_name)

        # Oluşturma tarihi bilgisi
        created_at = datetime.fromtimestamp(os.path.getctime(folder_path)).isoformat()

        return jsonify({
            'success': True,
            'kubik_name': kubik_name,
            'files': files,
            'created_at': created_at,
            'file_count': len(files)
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/kubik/create', methods=['POST'])
def create_kubik():
    """Yeni kübik oluştur"""
    try:
        kubik_name = request.form.get('kubik_name')
        if not kubik_name:
            return jsonify({'success': False, 'message': 'Kübik adı gerekli'})

        if 'files' not in request.files:
            return jsonify({'success': False, 'message': 'Dosya yüklenmedi'})

        files = request.files.getlist('files')
        if not files or all(file.filename == '' for file in files):
            return jsonify({'success': False, 'message': 'Geçerli dosya seçilmedi'})

        # Kübik klasörünü oluştur
        folder_path = os.path.join(DATA_FOLDER, kubik_name)
        os.makedirs(folder_path, exist_ok=True)

        # Dosyaları kaydet
        saved_files = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(folder_path, filename)
                file.save(file_path)
                saved_files.append(filename)

        if not saved_files:
            return jsonify({'success': False, 'message': 'Geçerli dosya bulunamadı (PDF, DOCX desteklenir)'})

        # Belgeleri işle
        docs = load_and_process_documents(folder_path)
        if not docs:
            return jsonify({'success': False, 'message': 'Dosyalar işlenemedi'})

        # Vektör veritabanı oluştur
        index_path = os.path.join(folder_path, "faiss_index")
        vector_store = create_or_load_vector_store(docs, index_path)
        if not vector_store:
            return jsonify({'success': False, 'message': 'Vektör veritabanı oluşturulamadı'})

        # Chatbot'u hazırla
        retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5, "fetch_k": 10})

        # Thread-safe LLM kullan
        llm = ThreadSafeLLM(google_api_key, temperature=0.2)

        # Aktif chatbot'lara ekle
        chatbot_id = str(uuid.uuid4())
        active_chatbots[chatbot_id] = {
            'kubik_name': kubik_name,
            'retriever': retriever,
            'llm': llm,
            'created_at': datetime.now().isoformat()
        }

        # Geçmişe ekle
        add_kubik_to_history(kubik_name)

        return jsonify({
            'success': True,
            'message': f'Kübik "{kubik_name}" başarıyla oluşturuldu',
            'chatbot_id': chatbot_id,
            'kubik_name': kubik_name,
            'files': saved_files
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/kubik/load/<kubik_name>', methods=['POST'])
def load_kubik(kubik_name):
    """Mevcut kübiği yükle"""
    try:
        folder_path = os.path.join(DATA_FOLDER, kubik_name)
        index_path = os.path.join(folder_path, "faiss_index")

        if not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': f'Kübik "{kubik_name}" bulunamadı'})

        if not os.path.exists(index_path):
            return jsonify({'success': False, 'message': f'Kübik "{kubik_name}" vektör veritabanı bulunamadı'})

        # Vektör veritabanını yükle
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        vector_store = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5, "fetch_k": 10})

        # Thread-safe LLM kullan
        llm = ThreadSafeLLM(google_api_key, temperature=0.5)

        # Aktif chatbot'lara ekle
        chatbot_id = str(uuid.uuid4())
        active_chatbots[chatbot_id] = {
            'kubik_name': kubik_name,
            'retriever': retriever,
            'llm': llm,
            'created_at': datetime.now().isoformat()
        }

        return jsonify({
            'success': True,
            'message': f'Kübik "{kubik_name}" başarıyla yüklendi',
            'chatbot_id': chatbot_id,
            'kubik_name': kubik_name
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/kubik/update', methods=['POST'])
def update_kubik():
    """Mevcut kübiği güncelle - yeni dosyalar ekle"""
    try:
        kubik_name = request.form.get('kubik_name')
        if not kubik_name:
            return jsonify({'success': False, 'message': 'Kübik adı gerekli'})

        folder_path = os.path.join(DATA_FOLDER, kubik_name)
        if not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': f'Kübik "{kubik_name}" bulunamadı'})

        if 'files' not in request.files:
            return jsonify({'success': False, 'message': 'Yeni dosya yüklenmedi'})

        files = request.files.getlist('files')
        if not files or all(file.filename == '' for file in files):
            return jsonify({'success': False, 'message': 'Geçerli yeni dosya seçilmedi'})

        # Yeni dosyaları kaydet
        saved_files = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(folder_path, filename)

                # Aynı isimli dosya varsa üzerine yaz
                if os.path.exists(file_path):
                    os.remove(file_path)

                file.save(file_path)
                saved_files.append(filename)

        if not saved_files:
            return jsonify({'success': False, 'message': 'Geçerli dosya bulunamadı (PDF, DOCX desteklenir)'})

        # Tüm belgeleri yeniden işle
        docs = load_and_process_documents(folder_path)
        if not docs:
            return jsonify({'success': False, 'message': 'Dosyalar işlenemedi'})

        # Eski vektör veritabanını sil
        index_path = os.path.join(folder_path, "faiss_index")
        if os.path.exists(index_path):
            shutil.rmtree(index_path)

        # Yeni vektör veritabanı oluştur
        vector_store = create_or_load_vector_store(docs, index_path)
        if not vector_store:
            return jsonify({'success': False, 'message': 'Vektör veritabanı oluşturulamadı'})

        # Aktif chatbot'ları güncelle (varsa)
        for chatbot_id, chatbot_info in list(active_chatbots.items()):
            if chatbot_info['kubik_name'] == kubik_name:
                retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5, "fetch_k": 10})
                active_chatbots[chatbot_id]['retriever'] = retriever

        return jsonify({
            'success': True,
            'message': f'Kübik "{kubik_name}" başarıyla güncellendi',
            'added_files': saved_files,
            'total_files': len(get_kubik_files(kubik_name))
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/kubik/remove-file', methods=['POST'])
def remove_file_from_kubik():
    """Kübikten dosya sil"""
    try:
        data = request.get_json()
        kubik_name = data.get('kubik_name')
        file_name = data.get('file_name')

        if not kubik_name or not file_name:
            return jsonify({'success': False, 'message': 'Kübik adı ve dosya adı gerekli'})

        folder_path = os.path.join(DATA_FOLDER, kubik_name)
        file_path = os.path.join(folder_path, file_name)

        if not os.path.exists(file_path):
            return jsonify({'success': False, 'message': 'Dosya bulunamadı'})

        # Dosyayı sil
        os.remove(file_path)

        # Kalan dosyaları kontrol et
        remaining_files = get_kubik_files(kubik_name)

        if len(remaining_files) == 0:
            return jsonify({'success': False, 'message': 'Kübikteki son dosyayı silemezsiniz'})

        # Vektör veritabanını yeniden oluştur
        docs = load_and_process_documents(folder_path)
        if docs:
            index_path = os.path.join(folder_path, "faiss_index")
            if os.path.exists(index_path):
                shutil.rmtree(index_path)

            vector_store = create_or_load_vector_store(docs, index_path)

            # Aktif chatbot'ları güncelle (varsa)
            for chatbot_id, chatbot_info in list(active_chatbots.items()):
                if chatbot_info['kubik_name'] == kubik_name:
                    if vector_store:
                        retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5, "fetch_k": 10})
                        active_chatbots[chatbot_id]['retriever'] = retriever

        return jsonify({
            'success': True,
            'message': f'"{file_name}" dosyası silindi',
            'remaining_files': len(remaining_files)
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/kubik/delete', methods=['POST'])
def delete_kubik():
    """Kübik'i tamamen sil"""
    try:
        data = request.get_json()
        kubik_name = data.get('kubik_name')

        if not kubik_name:
            return jsonify({'success': False, 'message': 'Kübik adı gerekli'})

        folder_path = os.path.join(DATA_FOLDER, kubik_name)
        if not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Kübik bulunamadı'})

        # Aktif chatbot'ları temizle
        chatbots_to_remove = []
        for chatbot_id, chatbot_info in active_chatbots.items():
            if chatbot_info['kubik_name'] == kubik_name:
                chatbots_to_remove.append(chatbot_id)

        for chatbot_id in chatbots_to_remove:
            del active_chatbots[chatbot_id]

        # Klasörü sil
        shutil.rmtree(folder_path)

        # Geçmişten kaldır
        remove_kubik_from_history(kubik_name)

        return jsonify({
            'success': True,
            'message': f'Kübik "{kubik_name}" başarıyla silindi'
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/chat/stream', methods=['POST'])
def chat_stream():
    """Streaming chat endpoint"""
    try:
        data = request.get_json()
        chatbot_id = data.get('chatbot_id')
        message = data.get('message')
        history = data.get('history', [])

        if not chatbot_id or chatbot_id not in active_chatbots:
            return jsonify({'success': False, 'message': 'Geçersiz chatbot ID'})

        if not message:
            return jsonify({'success': False, 'message': 'Mesaj boş'})

        chatbot = active_chatbots[chatbot_id]
        retriever = chatbot['retriever']
        llm = chatbot['llm']

        def generate():
            try:
                # İlgili belgeleri al
                docs = retriever.invoke(message)
                knowledge = "\n\n".join([doc.page_content for doc in docs])

                # Geçmişi formatla
                history_str = ""
                for msg in history:
                    role = "Kullanıcı" if msg['role'] == 'user' else "Bot"
                    history_str += f"{role}: {msg['content']}\n"

                # RAG prompt oluştur
                rag_prompt = f"""
Sen bir belge analiz uzmanısın. Aşağıdaki bilgiye dayanarak soruyu yanıtla.
Sadece verilen bilgiye dayanarak yanıt ver. Bilgi yoksa "Üzgünüm, bu sorunun cevabını belgelerde bulamadım." de.

--- SORU ---
{message}

--- GEÇMİŞ ---
{history_str}

--- BİLGİ ---
{knowledge}
"""

                # Streaming response
                full_response = ""
                for chunk in llm.stream(rag_prompt):
                    if chunk.content:
                        full_response += chunk.content
                        yield f"data: {json.dumps({'content': chunk.content, 'done': False})}\n\n"

                yield f"data: {json.dumps({'content': '', 'done': True, 'full_response': full_response})}\n\n"

            except Exception as e:
                yield f"data: {json.dumps({'error': str(e)})}\n\n"

        return Response(generate(), mimetype='text/plain')

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/chat/send', methods=['POST'])
def chat_send():
    """Non-streaming chat endpoint"""
    try:
        data = request.get_json()
        chatbot_id = data.get('chatbot_id')
        message = data.get('message')
        history = data.get('history', [])

        if not chatbot_id or chatbot_id not in active_chatbots:
            return jsonify({'success': False, 'message': 'Geçersiz chatbot ID'})

        if not message:
            return jsonify({'success': False, 'message': 'Mesaj boş'})

        chatbot = active_chatbots[chatbot_id]
        retriever = chatbot['retriever']
        llm = chatbot['llm']

        # İlgili belgeleri al
        docs = retriever.invoke(message)
        knowledge = "\n\n".join([doc.page_content for doc in docs])

        # Geçmişi formatla
        history_str = ""
        for msg in history:
            role = "Kullanıcı" if msg['role'] == 'user' else "Bot"
            history_str += f"{role}: {msg['content']}\n"

        # RAG prompt oluştur
        rag_prompt = f"""
Sen bir belge analiz uzmanısın. Aşağıdaki bilgiye dayanarak soruyu yanıtla.
Sadece verilen bilgiye dayanarak yanıt ver. Bilgi yoksa "Üzgünüm, bu sorunun cevabını belgelerde bulamadım." de.

--- SORU ---
{message}

--- GEÇMİŞ ---
{history_str}

--- BİLGİ ---
{knowledge}
"""

        # Response al
        response = llm.invoke(rag_prompt)

        return jsonify({
            'success': True,
            'response': response.content
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/chatbot/info/<chatbot_id>', methods=['GET'])
def get_chatbot_info(chatbot_id):
    """Chatbot bilgilerini getir"""
    if chatbot_id not in active_chatbots:
        return jsonify({'success': False, 'message': 'Chatbot bulunamadı'})

    chatbot = active_chatbots[chatbot_id]
    return jsonify({
        'success': True,
        'kubik_name': chatbot['kubik_name'],
        'created_at': chatbot['created_at']
    })


@app.route('/api/chatbot/close/<chatbot_id>', methods=['POST'])
def close_chatbot(chatbot_id):
    """Chatbot'u kapat"""
    if chatbot_id in active_chatbots:
        del active_chatbots[chatbot_id]
        return jsonify({'success': True, 'message': 'Chatbot kapatıldı'})
    return jsonify({'success': False, 'message': 'Chatbot bulunamadı'})


@app.route('/')
def index():
    """Ana sayfa - frontend'i serve et"""
    return send_from_directory(STATIC_FOLDER, 'index.html')


@app.route('/static/<path:filename>')
def static_files(filename):
    """Static dosyaları serve et"""
    return send_from_directory(STATIC_FOLDER, filename)


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)