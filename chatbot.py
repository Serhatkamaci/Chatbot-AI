import os
import shutil
import json
from dotenv import load_dotenv
import gradio as gr
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from docx import Document
from langchain.schema import Document as LangchainDocument
from zipfile import ZipFile

load_dotenv()
google_api_key = os.getenv("GOOGLE_API_KEY")

if not google_api_key:
    raise ValueError("GOOGLE_API_KEY ortam değişkeni ayarlanmamış.")

# --- Kübik geçmişini yönetme ---
HISTORY_FILE = "kubik_history.json"


def load_kubik_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []


def save_kubik_history(history):
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def add_kubik_to_history(kubik_name):
    history = load_kubik_history()
    # Eğer zaten varsa, en üste taşı
    history = [h for h in history if h != kubik_name]
    history.insert(0, kubik_name)
    # En fazla 20 kübik tut
    if len(history) > 20:
        history = history[:20]
    save_kubik_history(history)
    return history


# --- DOCX dosyasını güvenli açma ve metin alma ---
def load_docx_with_python_docx(file_path):
    try:
        doc = Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip() != ""]
        if not paragraphs:
            print(f"DOCX dosyası içerik olarak boş: {file_path}")
            return []
        full_text = "\n".join(paragraphs)
        return [LangchainDocument(page_content=full_text, metadata={"source": file_path})]
    except Exception as e:
        print(f"DOCX dosyası açılamadı ({file_path}): {e}")
        return []


# --- Belgeleri Yükle ve Böl ---
def load_and_process_documents(folder_path):
    documents = []
    pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".pdf")]
    docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]

    if not pdf_files and not docx_files:
        print(f"'{folder_path}' klasöründe PDF veya DOCX dosyası bulunamadı.")
        return None

    for pdf_file in pdf_files:
        file_path = os.path.join(folder_path, pdf_file)
        print(f"'{pdf_file}' yükleniyor...")
        loader = PyPDFLoader(file_path)
        documents.extend(loader.load())

    for docx_file in docx_files:
        file_path = os.path.join(folder_path, docx_file)
        print(f"'{docx_file}' yükleniyor...")
        documents.extend(load_docx_with_python_docx(file_path))

    if not documents:
        return None

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=300,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    return text_splitter.split_documents(documents)


# --- Vektör Veritabanı Yarat veya Yükle ---
def create_or_load_vector_store(docs, index_path):
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=google_api_key
    )

    if os.path.exists(index_path):
        print("Vektör veritabanı yükleniyor...")
        return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)

    if not docs:
        print("Vektör veritabanı oluşturmak için belge yok.")
        return None

    print("Yeni vektör veritabanı oluşturuluyor...")
    vector_store = FAISS.from_documents(docs, embeddings)
    vector_store.save_local(index_path)
    return vector_store


# --- Global değişkenler ---
retriever = None
llm = None
current_kubik_name = ""


# --- Stream Response ---
def stream_response(message, history):
    if not message:
        yield "Lütfen bir soru girin."
        return

    global retriever, llm

    if not retriever or not llm:
        yield "Chatbot hazır değil. Lütfen önce dosyaları yükleyip chatbot oluşturun."
        return

    docs = retriever.invoke(message)
    knowledge = "\n\n".join([doc.page_content for doc in docs])

    rag_prompt = f"""
Sen bir belge analiz uzmanısın. Aşağıdaki bilgiye dayanarak soruyu yanıtla.
Sadece verilen bilgiye dayanarak yanıt ver. Bilgi yoksa "Üzgünüm, bu sorunun cevabını belgelerde bulamadım." de.

--- SORU ---
{message}

--- GEÇMİŞ ---
{history}

--- BİLGİ ---
{knowledge}
"""

    partial_message = ""
    for response in llm.stream(rag_prompt):
        partial_message += response.content
        yield partial_message


# --- Mevcut kübiği yükle ---
def load_existing_kubik(kubik_name):
    global retriever, llm, current_kubik_name

    if not kubik_name:
        return "Kübik adı boş.", False

    folder_path = os.path.join("data", kubik_name)
    index_path = os.path.join(folder_path, "faiss_index")

    if not os.path.exists(folder_path):
        return f"'{kubik_name}' kübik klasörü bulunamadı.", False

    if not os.path.exists(index_path):
        return f"'{kubik_name}' kübik vektör veritabanı bulunamadı.", False

    try:
        # Vektör veritabanını yükle
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        vector_store = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5, "fetch_k": 10})

        # LLM'yi başlat
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash-latest",
            google_api_key=google_api_key,
            temperature=0.5,
        )

        current_kubik_name = kubik_name
        return f"'{kubik_name}' kübik başarıyla yüklendi!", True

    except Exception as e:
        return f"'{kubik_name}' kübik yüklenirken hata: {str(e)}", False


# --- Dosyaları kaydet ve chatbot hazırla ---
def save_files_and_prepare_chat(kubik_name, files):
    global current_kubik_name

    if not kubik_name:
        return "Lütfen geçerli bir Kübik Adı girin.", False, []

    if not files:
        return "Lütfen en az bir dosya yükleyin.", False, []

    folder_path = os.path.join("data", kubik_name)
    os.makedirs(folder_path, exist_ok=True)

    for file in files:
        # Gradio'da yüklenen dosya objesi
        if hasattr(file, 'name') and file.name:
            # file.name dosya yolunu içerir
            original_filename = os.path.basename(file.name)
            file_path = file.name
        else:
            return "Dosya bilgisi alınamadı.", False, []

        # Sadece pdf ve docx kabul et
        if not (original_filename.lower().endswith(".pdf") or original_filename.lower().endswith(".docx")):
            return f"Sadece PDF ve DOCX dosyaları kabul edilir: {original_filename}", False, []

        save_path = os.path.join(folder_path, original_filename)

        # Dosyayı kopyala (Gradio geçici dosyasından hedef konuma)
        try:
            shutil.copy2(file_path, save_path)
            file_size = os.path.getsize(save_path)
            print(f"Dosya kopyalandı: {save_path}, boyut: {file_size} bytes")

            # Dosya boyutunu kontrol et
            if file_size < 1000:  # 1KB'den küçükse sorun var
                print(f"UYARI: Dosya çok küçük görünüyor: {original_filename} ({file_size} bytes)")

        except Exception as e:
            return f"Dosya kopyalanırken hata: {e}", False, []

    # Belgeleri yükle
    docs = load_and_process_documents(folder_path)
    if not docs:
        return f"'{kubik_name}' klasöründe uygun dosya bulunamadı veya dosyalar bozuk.", False, []

    index_path = os.path.join(folder_path, "faiss_index")

    global retriever, llm
    vector_store = create_or_load_vector_store(docs, index_path)
    if not vector_store:
        return "Vektör veritabanı oluşturulamadı.", False, []

    retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5, "fetch_k": 10})

    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash-latest",
        google_api_key=google_api_key,
        temperature=0.5,
    )

    current_kubik_name = kubik_name

    # Kübik geçmişine ekle
    history = add_kubik_to_history(kubik_name)

    return f"'{kubik_name}' chatbot'u başarıyla oluşturuldu! Artık sorularınızı sorabilirsiniz.", True, history


# --- Sidebar toggle fonksiyonu ---
def toggle_sidebar(current_visible):
    return gr.update(visible=not current_visible)


# --- Gradio UI ---
with gr.Blocks(title="Kübik Chatbot", css="""
.sidebar-button {
    position: fixed !important;
    top: 20px !important;
    left: 20px !important;
    z-index: 1000 !important;
    border-radius: 50% !important;
    width: 50px !important;
    height: 50px !important;
    padding: 0 !important;
}
.sidebar-container {
    position: fixed !important;
    top: 0 !important;
    left: 0 !important;
    width: 300px !important;
    height: 100vh !important;
    background: white !important;
    border-right: 1px solid #ddd !important;
    z-index: 999 !important;
    overflow-y: auto !important;
    padding: 70px 15px 15px 15px !important;
    box-shadow: 2px 0 5px rgba(0,0,0,0.1) !important;
}
.main-content {
    margin-left: 0px !important;
    transition: margin-left 0.3s ease !important;
}
.main-content.sidebar-open {
    margin-left: 300px !important;
}
""") as demo:
    sidebar_visible = gr.State(value=False)  # <- GÖRÜNÜRLÜK DURUMU TUTULUYOR

    sidebar_toggle = gr.Button("☰", elem_classes="sidebar-button", variant="secondary")

    with gr.Column(visible=False, elem_classes="sidebar-container") as sidebar:
        ...

    sidebar_toggle.click(
        toggle_sidebar,
        inputs=[sidebar_visible],
        outputs=[sidebar, sidebar_visible]
    )

    # Sidebar toggle butonu


    # Sidebar paneli
    with gr.Column(visible=False, elem_classes="sidebar-container") as sidebar:
        gr.Markdown("### 📚 Kübik Geçmişi")

        history_list = gr.Radio(
            choices=[],
            label="Mevcut Kübikler",
            interactive=True,
            value=None
        )

        load_kubik_btn = gr.Button("📂 Kübiği Yükle", variant="secondary", size="sm")

        gr.Markdown("---")

        sidebar_new_btn = gr.Button(
            "➕ Yeni Kübik Oluştur",
            variant="primary",
            size="sm"
        )

    # Ana içerik alanı
    with gr.Column(elem_classes="main-content") as main_content:
        # Ana form alanları (başlangıçta görünür)
        with gr.Group(visible=True) as setup_section:
            gr.Markdown("# 🚀 Kübik Dosya Yükleyici ve Chatbot Oluşturucu")
            gr.Markdown("### Belgelerinizi yükleyin ve özel chatbot'unuzu oluşturun")

            with gr.Row():
                kubik_name = gr.Textbox(
                    label="Kübik Adı (Klasör ismi)",
                    placeholder="Örn: project1",
                    scale=2
                )

            file_upload = gr.File(
                label="📁 Dosyalarınızı sürükleyip bırakın (PDF, DOCX)",
                file_types=[".pdf", ".docx"],
                file_count="multiple",
                height=120
            )

            with gr.Row():
                create_btn = gr.Button(
                    "🤖 Chatbot Oluştur",
                    variant="primary",
                    size="lg",
                    scale=1
                )

            output_message = gr.Textbox(
                label="📋 Durum",
                interactive=False,
                lines=2
            )

        # Chatbot arayüzü (başlangıçta gizli)
        with gr.Group(visible=False) as chat_section:
            with gr.Row():
                current_kubik_display = gr.Markdown("# 💬 Chatbot - Sorularınızı Sorun")
                new_chatbot_btn = gr.Button(
                    "🔄 Yeni Kübik Oluştur",
                    variant="secondary",
                    size="sm"
                )

            chatbot = gr.Chatbot(
                type="messages",
                height=650,
                show_label=False,
                container=True,
                bubble_full_width=False,
                show_copy_button=True
            )

            with gr.Row():
                user_input = gr.Textbox(
                    placeholder="💭 Sorularınızı buraya yazın...",
                    lines=2,
                    label="Mesajınız",
                    scale=4,
                    container=True
                )
                send_btn = gr.Button("📤 Gönder", variant="primary", scale=1)


    def on_create_click(kubik_name, files):
        msg, success, history = save_files_and_prepare_chat(kubik_name, files)
        if success:
            return (
                msg,  # output_message
                gr.update(visible=False),  # setup_section gizle
                gr.update(visible=True),  # chat_section göster
                [],  # chatbot geçmişini temizle
                gr.update(choices=history, value=None),  # history_list güncelle
                f"# 💬 {kubik_name} - Sorularınızı Sorun"  # current_kubik_display
            )
        else:
            history = load_kubik_history()
            return (
                msg,  # output_message
                gr.update(visible=True),  # setup_section görünür kalsın
                gr.update(visible=False),  # chat_section gizli kalsın
                [],  # chatbot
                gr.update(choices=history, value=None),  # history_list güncelle
                "# 💬 Chatbot - Sorularınızı Sorun"  # current_kubik_display
            )


    def reset_to_setup():
        global retriever, llm, current_kubik_name
        retriever = None
        llm = None
        current_kubik_name = ""
        history = load_kubik_history()
        return (
            gr.update(visible=True),  # setup_section göster
            gr.update(visible=False),  # chat_section gizle
            "",  # kubik_name temizle
            None,  # file_upload temizle
            "",  # output_message temizle
            [],  # chatbot geçmişi temizle
            "",  # user_input temizle
            gr.update(choices=history, value=None),  # history_list güncelle
            "# 💬 Chatbot - Sorularınızı Sorun"  # current_kubik_display
        )


    def load_kubik_from_history(selected_kubik):
        if not selected_kubik:
            return (
                "Lütfen bir kübik seçin.",  # output_message
                gr.update(visible=True),  # setup_section
                gr.update(visible=False),  # chat_section
                [],  # chatbot
                "# 💬 Chatbot - Sorularınızı Sorun"  # current_kubik_display
            )

        msg, success = load_existing_kubik(selected_kubik)
        if success:
            return (
                msg,  # output_message
                gr.update(visible=False),  # setup_section gizle
                gr.update(visible=True),  # chat_section göster
                [],  # chatbot geçmişini temizle
                f"# 💬 {selected_kubik} - Sorularınızı Sorun"  # current_kubik_display
            )
        else:
            return (
                msg,  # output_message
                gr.update(visible=True),  # setup_section
                gr.update(visible=False),  # chat_section
                [],  # chatbot
                "# 💬 Chatbot - Sorularınızı Sorun"  # current_kubik_display
            )


    def chatbot_interact(user_message, chat_history):
        if not user_message.strip():
            return chat_history, ""

        history_str = ""
        if chat_history:
            history_pairs = []
            for i in range(0, len(chat_history), 2):
                if i + 1 < len(chat_history):
                    user_msg = chat_history[i]['content'] if isinstance(chat_history[i], dict) else str(chat_history[i])
                    bot_msg = chat_history[i + 1]['content'] if isinstance(chat_history[i + 1], dict) else str(
                        chat_history[i + 1])
                    history_pairs.append(f"Kullanıcı: {user_msg}\nBot: {bot_msg}")
            history_str = "\n\n".join(history_pairs)

        bot_response = ""
        new_history = chat_history + [{"role": "user", "content": user_message}]

        for partial in stream_response(user_message, history_str):
            bot_response = partial
            # Son mesajı güncelle veya yeni bot mesajı ekle
            if len(new_history) > 0 and new_history[-1].get('role') == 'assistant':
                new_history[-1]['content'] = bot_response
            else:
                new_history.append({"role": "assistant", "content": bot_response})
            yield new_history, ""


    # Component başlangıç durumu
    demo.load(
        lambda: gr.update(choices=load_kubik_history(), value=None),
        outputs=[history_list]
    )

    # Event handlers
    sidebar_toggle.click(
        toggle_sidebar,
        inputs=[sidebar_visible],
        outputs=[sidebar, sidebar_visible]
    )

    create_btn.click(
        fn=on_create_click,
        inputs=[kubik_name, file_upload],
        outputs=[output_message, setup_section, chat_section, chatbot, history_list, current_kubik_display],
    )

    new_chatbot_btn.click(
        fn=reset_to_setup,
        inputs=[],
        outputs=[setup_section, chat_section, kubik_name, file_upload, output_message, chatbot, user_input,
                 history_list, current_kubik_display],
    )

    sidebar_new_btn.click(
        fn=reset_to_setup,
        inputs=[],
        outputs=[setup_section, chat_section, kubik_name, file_upload, output_message, chatbot, user_input,
                 history_list, current_kubik_display],
    )

    load_kubik_btn.click(
        fn=load_kubik_from_history,
        inputs=[history_list],
        outputs=[output_message, setup_section, chat_section, chatbot, current_kubik_display],
    )

    user_input.submit(
        chatbot_interact,
        inputs=[user_input, chatbot],
        outputs=[chatbot, user_input],
    )

    send_btn.click(
        chatbot_interact,
        inputs=[user_input, chatbot],
        outputs=[chatbot, user_input],
    )

demo.launch()